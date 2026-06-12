from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "final_report" / "customer_churn_report.docx"
FIG = ROOT / "reports" / "figures"
MET = ROOT / "models" / "metrics"


def load_data():
    results_df = pd.read_csv(MET / "model_results.csv")
    best_summary = pd.read_csv(MET / "best_model_summary.csv").iloc[0]
    feature_cmp = pd.read_csv(MET / "feature_selection_comparison.csv")
    selected_features = pd.read_csv(MET / "selected_features.csv")["feature"].tolist()
    return results_df, best_summary, feature_cmp, selected_features


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


class ReportBuilder:
    def __init__(self):
        self.doc = Document()
        self.figures = [
            ("Hình 1", "Phân bố khách hàng theo trạng thái Churn", FIG / "churn_distribution.png"),
            ("Hình 2", "Tỷ lệ Churn theo loại hợp đồng", FIG / "contract_vs_churn.png"),
            ("Hình 3", "Phân bố MonthlyCharges theo Churn", FIG / "monthly_charges_by_churn.png"),
            ("Hình 4", "So sánh dữ liệu trước và sau SMOTENC", FIG / "so_sanh_truoc_sau_smote.png"),
            ("Hình 5", "So sánh hiệu suất các mô hình", FIG / "bieudocot_sosanh_model.png"),
            ("Hình 6", "Confusion Matrix của mô hình tốt nhất", FIG / "xgboost_confusion_matrix_step06.png"),
            ("Hình 7", "ROC Curve của mô hình tốt nhất", FIG / "xgboost_roc_curve_step06.png"),
            ("Hình 8", "Feature Importance của mô hình tốt nhất", FIG / "xgboost_feature_importance_step06.png"),
        ]
        self.figures = [item for item in self.figures if item[2].exists()]
        self.setup_styles()

    def setup_styles(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.10

        for name, size, color, before, after in [
            ("Heading 1", 16, "2E74B5", 16, 8),
            ("Heading 2", 13, "2E74B5", 12, 6),
            ("Heading 3", 12, "1F4D78", 8, 4),
        ]:
            style = styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        caption = styles.add_style("Caption Custom", 1)
        caption.font.name = "Calibri"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        caption.font.size = Pt(10)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor(89, 89, 89)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)

    def add_center(self, text: str, size=12, bold=False, color=None, after=6):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        r.font.size = Pt(size)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        return p

    def add_para(self, text: str):
        self.doc.add_paragraph(text)

    def add_bullets(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(str(item))

    def add_numbers(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(str(item))

    def add_table(self, headers, rows, widths=None, font_size=9):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            set_cell_shading(table.rows[0].cells[i], "F2F4F7")
            set_cell_text(
                table.rows[0].cells[i],
                header,
                bold=True,
                size=font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                text = str(value)
                is_number = text.replace(".", "", 1).replace("-", "", 1).isdigit()
                align = WD_ALIGN_PARAGRAPH.CENTER if is_number else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_text(cells[i], text, size=font_size, align=align)
        if widths:
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = Inches(width)
        self.doc.add_paragraph()
        return table

    def add_figure(self, label: str, caption: str, path: Path, width=6.0):
        if not path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        self.doc.add_paragraph(f"{label}. {caption}", style="Caption Custom")

    def build(self):
        results_df, best_summary, feature_cmp, selected_features = load_data()

        # Cover page
        self.add_center("BỘ GIÁO DỤC VÀ ĐÀO TẠO", 13, True)
        self.add_center("ĐẠI HỌC PHENIKAA", 13, True, after=28)
        self.add_center("BÁO CÁO HỌC PHẦN", 18, True, "0B2545")
        self.add_center("LẬP TRÌNH PHÂN TÍCH DỮ LIỆU VỚI PYTHON", 15, True, "0B2545", after=24)
        self.add_center("Đề tài:", 13, True)
        self.add_center(
            "PHÂN TÍCH VÀ DỰ ĐOÁN KHÁCH HÀNG RỜI BỎ DỊCH VỤ VIỄN THÔNG",
            16,
            True,
            "2E74B5",
            after=30,
        )
        self.add_center("Customer Churn Prediction", 13, False, "555555", after=32)
        self.add_center("Thành viên nhóm: ........................................................", 12)
        self.add_center("Lớp: ........................................................", 12)
        self.add_center("Giảng viên hướng dẫn: ........................................................", 12, after=80)
        self.add_center("Hà Nội - Năm 2026", 12, True)
        self.doc.add_page_break()

        self.doc.add_heading("PHÂN CÔNG NHIỆM VỤ", level=1)
        self.add_table(
            ["STT", "Thành viên", "Nhiệm vụ", "Mức độ hoàn thành"],
            [
                ["1", "................................", "Thu thập dữ liệu, tìm hiểu dataset, viết phần giới thiệu", "100%"],
                ["2", "................................", "Tiền xử lý dữ liệu, xử lý missing values, encode, scale, SMOTENC", "100%"],
                ["3", "................................", "EDA, trực quan hóa dữ liệu, phân tích insight", "100%"],
                ["4", "................................", "Huấn luyện mô hình, đánh giá, viết báo cáo và slide", "100%"],
            ],
            widths=[0.45, 1.7, 3.6, 1.0],
            font_size=9,
        )
        self.doc.add_page_break()

        self.doc.add_heading("MỤC LỤC", level=1)
        self.add_bullets(
            [
                "1. Mở đầu",
                "2. Tóm tắt",
                "3. Giới thiệu đề tài",
                "4. Dữ liệu sử dụng",
                "5. Kiến thức nền tảng",
                "6. Quy trình thực hiện theo 7 bước",
                "7. Phân tích khám phá dữ liệu",
                "8. Tiền xử lý dữ liệu",
                "9. Lựa chọn feature",
                "10. Huấn luyện và đánh giá mô hình",
                "11. Nhận xét và ý nghĩa thực tiễn",
                "12. Kết luận",
            ]
        )
        self.doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
        self.add_bullets([f"{label}. {caption}" for label, caption, _ in self.figures])
        self.doc.add_heading("DANH MỤC BẢNG", level=1)
        self.add_bullets(
            [
                "Bảng 1. Các nhóm thuộc tính trong bộ dữ liệu",
                "Bảng 2. Danh sách feature được chọn bằng SelectKBest",
                "Bảng 3. So sánh hiệu quả full features và selected features",
                "Bảng 4. Kết quả đánh giá các mô hình",
                "Bảng 5. Tóm tắt mô hình tốt nhất",
            ]
        )
        self.doc.add_page_break()

        self.doc.add_heading("MỞ ĐẦU", level=1)
        self.add_para(
            "Trong bối cảnh dữ liệu ngày càng giữ vai trò quan trọng trong hoạt động kinh doanh, "
            "việc khai thác dữ liệu khách hàng giúp doanh nghiệp hiểu rõ hành vi sử dụng dịch vụ "
            "và đưa ra quyết định chính xác hơn. Đối với doanh nghiệp viễn thông, khách hàng rời bỏ "
            "dịch vụ là một vấn đề quan trọng vì chi phí giữ chân khách hàng cũ thường thấp hơn chi phí "
            "thu hút khách hàng mới."
        )
        self.add_para(
            "Xuất phát từ thực tiễn đó, nhóm lựa chọn đề tài “Phân tích và dự đoán khách hàng rời bỏ dịch vụ viễn thông”. "
            "Đề tài giúp áp dụng kiến thức phân tích dữ liệu với Python và có ý nghĩa thực tế trong việc hỗ trợ doanh nghiệp "
            "nhận diện sớm nhóm khách hàng có nguy cơ rời bỏ."
        )

        self.doc.add_heading("TÓM TẮT", level=1)
        self.add_para(
            "Báo cáo trình bày quá trình xây dựng hệ thống dự đoán khách hàng rời bỏ dịch vụ viễn thông "
            "dựa trên bộ dữ liệu Telco Customer Churn. Bài toán được xác định là phân loại nhị phân với biến mục tiêu là Churn."
        )
        self.add_para(
            "Nhóm thực hiện đầy đủ quy trình phân tích dữ liệu gồm thu thập dữ liệu, làm sạch dữ liệu, EDA, "
            "tiền xử lý, xử lý mất cân bằng bằng SMOTENC, chuẩn hóa dữ liệu bằng StandardScaler, huấn luyện nhiều mô hình "
            "và đánh giá bằng nhiều chỉ số."
        )
        self.add_para(
            f"Kết quả thực nghiệm cho thấy mô hình {best_summary['best_model']} đạt F1-score "
            f"{best_summary['f1']:.4f} và ROC-AUC {best_summary['roc_auc']:.4f}. Đây là mô hình tốt nhất "
            "theo tiêu chí F1-score trong lần chạy hiện tại."
        )

        self.doc.add_heading("GIỚI THIỆU ĐỀ TÀI", level=1)
        self.add_para(
            "Customer Churn là thuật ngữ chỉ việc khách hàng ngừng sử dụng sản phẩm hoặc dịch vụ của doanh nghiệp. "
            "Trong ngành viễn thông, churn có thể xảy ra khi khách hàng chuyển sang nhà mạng khác, hủy hợp đồng "
            "hoặc không tiếp tục sử dụng dịch vụ."
        )
        self.add_para(
            "Mục tiêu của đề tài là dự đoán khách hàng nào có khả năng churn dựa trên các đặc điểm như thời gian sử dụng dịch vụ, "
            "loại hợp đồng, chi phí hàng tháng, dịch vụ Internet, hỗ trợ kỹ thuật và phương thức thanh toán."
        )
        self.add_bullets(
            [
                "Dữ liệu thực tế, dễ hiểu và phù hợp với bài toán phân tích dữ liệu.",
                "Có đầy đủ các bước từ làm sạch dữ liệu đến trực quan hóa và machine learning.",
                "Có thể so sánh nhiều mô hình phân loại khác nhau.",
                "Có ý nghĩa thực tiễn trong việc hỗ trợ doanh nghiệp giữ chân khách hàng.",
            ]
        )

        self.doc.add_heading("DỮ LIỆU SỬ DỤNG", level=1)
        self.add_para(
            "Bộ dữ liệu được sử dụng là Telco Customer Churn Dataset. Mỗi dòng dữ liệu tương ứng với một khách hàng, "
            "mỗi cột tương ứng với một thuộc tính hoặc thông tin sử dụng dịch vụ của khách hàng."
        )
        self.add_table(
            ["Nhóm thuộc tính", "Các cột tiêu biểu", "Ý nghĩa"],
            [
                ["Thông tin cá nhân", "gender, SeniorCitizen, Partner, Dependents", "Mô tả đặc điểm cơ bản của khách hàng"],
                ["Thời gian và chi phí", "tenure, MonthlyCharges, TotalCharges", "Phản ánh mức độ gắn bó và chi phí sử dụng dịch vụ"],
                ["Dịch vụ", "InternetService, OnlineSecurity, TechSupport, StreamingTV", "Mô tả các dịch vụ khách hàng đang sử dụng"],
                ["Hợp đồng và thanh toán", "Contract, PaperlessBilling, PaymentMethod", "Mô tả hình thức hợp đồng và thanh toán"],
                ["Biến mục tiêu", "Churn", "Cho biết khách hàng có rời bỏ dịch vụ hay không"],
            ],
            widths=[1.4, 2.4, 2.6],
            font_size=9,
        )

        self.doc.add_heading("KIẾN THỨC NỀN TẢNG", level=1)
        self.doc.add_heading("Ngôn ngữ và thư viện sử dụng", level=2)
        self.add_bullets(
            [
                "Pandas: đọc, xử lý và phân tích dữ liệu dạng bảng.",
                "NumPy: hỗ trợ tính toán số học và xử lý mảng.",
                "Matplotlib và Seaborn: trực quan hóa dữ liệu bằng biểu đồ.",
                "Scikit-learn: cung cấp công cụ tiền xử lý, huấn luyện và đánh giá mô hình.",
                "Imbalanced-learn: hỗ trợ xử lý dữ liệu mất cân bằng bằng SMOTE/SMOTENC.",
                "XGBoost và LightGBM: các mô hình boosting mạnh cho dữ liệu dạng bảng.",
            ]
        )
        self.doc.add_heading("Các mô hình học máy", level=2)
        self.add_para(
            "Project sử dụng nhiều mô hình phân loại để so sánh khách quan, bao gồm Logistic Regression, Decision Tree, "
            "Random Forest, Gradient Boosting, AdaBoost, XGBoost, LightGBM và Voting Ensemble."
        )
        self.add_para(
            "Logistic Regression được dùng làm baseline vì đơn giản và dễ giải thích. Decision Tree mô hình hóa quyết định "
            "theo dạng cây. Random Forest kết hợp nhiều cây để tăng độ ổn định. Gradient Boosting, AdaBoost, XGBoost và "
            "LightGBM là các mô hình boosting, trong đó mô hình sau tập trung sửa lỗi của mô hình trước. Voting Ensemble "
            "kết hợp nhiều mô hình để đưa ra dự đoán tổng hợp."
        )

        self.doc.add_heading("QUY TRÌNH THỰC HIỆN THEO 7 BƯỚC", level=1)
        steps = [
            ("Bước 1 - Thu thập dữ liệu", "Đọc dữ liệu từ file CSV, kiểm tra kích thước, danh sách cột và một số dòng mẫu."),
            ("Bước 2 - Làm sạch và tiền xử lý dữ liệu", "Xử lý TotalCharges, loại bỏ customerID, encode dữ liệu phân loại, chia train/test, SMOTENC và StandardScaler."),
            ("Bước 3 - Khai phá và phân tích dữ liệu", "Thực hiện EDA bằng biểu đồ cột, tròn, histogram, boxplot và heatmap để tìm insight."),
            ("Bước 4 - Huấn luyện mô hình", "Huấn luyện nhiều mô hình phân loại, bao gồm các mô hình cơ bản, ensemble và boosting."),
            ("Bước 5 - Lựa chọn feature và tối ưu hóa nhẹ", "Dùng Random Forest importance, Mutual Information và SelectKBest để chọn feature quan trọng."),
            ("Bước 6 - Đánh giá và so sánh mô hình", "Đánh giá bằng Accuracy, Precision, Recall, F1-score và ROC-AUC, sau đó lưu best model."),
            ("Bước 7 - Truyền đạt kết quả", "Tổng hợp insight, kết luận và nội dung đưa vào báo cáo hoặc slide."),
        ]
        for title, desc in steps:
            self.doc.add_heading(title, level=2)
            self.add_para(desc)

        self.doc.add_heading("PHÂN TÍCH KHÁM PHÁ DỮ LIỆU", level=1)
        self.add_para(
            "Phần EDA giúp nhóm hiểu dữ liệu trước khi xây dựng mô hình. Các biểu đồ cho thấy churn có liên quan đến loại hợp đồng, "
            "thời gian sử dụng dịch vụ, chi phí hàng tháng và các dịch vụ hỗ trợ."
        )
        for label, caption, path in self.figures[:3]:
            self.add_figure(label, caption, path, width=5.7)
        self.add_para(
            "Từ các biểu đồ, có thể nhận thấy khách hàng dùng hợp đồng ngắn hạn và khách hàng có thời gian sử dụng dịch vụ thấp "
            "thường có khả năng churn cao hơn. MonthlyCharges cũng là một biến đáng chú ý vì chi phí hàng tháng cao có thể ảnh hưởng "
            "đến quyết định rời bỏ dịch vụ."
        )

        self.doc.add_heading("TIỀN XỬ LÝ DỮ LIỆU", level=1)
        self.add_para(
            "Trong bước tiền xử lý, nhóm làm sạch text, chuyển TotalCharges sang dạng số, điền giá trị thiếu bằng median, "
            "loại bỏ customerID, mã hóa dữ liệu phân loại, chia train/test, cân bằng dữ liệu bằng SMOTENC và chuẩn hóa các biến số."
        )
        self.add_figure("Hình 4", "So sánh dữ liệu trước và sau SMOTENC", FIG / "so_sanh_truoc_sau_smote.png", width=5.3)
        self.add_para(
            "SMOTENC được sử dụng vì dữ liệu có cả biến số và biến phân loại. Kỹ thuật này giúp giảm ảnh hưởng của mất cân bằng dữ liệu "
            "trong tập train, từ đó hỗ trợ mô hình học tốt hơn lớp khách hàng churn."
        )

        self.doc.add_heading("LỰA CHỌN FEATURE", level=1)
        self.add_para(
            "Bước lựa chọn feature nhằm xác định các biến quan trọng nhất đối với biến mục tiêu Churn. Nhóm sử dụng Random Forest "
            "feature importance, Mutual Information và SelectKBest."
        )
        self.add_table(
            ["STT", "Feature được chọn"],
            [[i + 1, feature] for i, feature in enumerate(selected_features)],
            widths=[0.7, 5.5],
            font_size=10,
        )
        feature_rows = [
            [
                row["model"],
                row["feature_set"],
                f"{row['accuracy']:.4f}",
                f"{row['recall']:.4f}",
                f"{row['f1']:.4f}",
                f"{row['roc_auc']:.4f}",
            ]
            for _, row in feature_cmp.iterrows()
        ]
        self.add_table(
            ["Model", "Feature set", "Accuracy", "Recall", "F1", "ROC-AUC"],
            feature_rows,
            widths=[1.35, 1.25, 0.85, 0.85, 0.75, 0.85],
            font_size=8,
        )
        self.add_para(
            "Kết quả cho thấy selected features giúp cải thiện F1-score ở một số mô hình, đặc biệt là Gradient Boosting và "
            "Logistic Regression. Vì vậy, ở bước đánh giá chính thức, Notebook 06 sử dụng chiến lược selected_features."
        )

        self.doc.add_heading("HUẤN LUYỆN VÀ ĐÁNH GIÁ MÔ HÌNH", level=1)
        self.add_para(
            "Các mô hình được đánh giá bằng Accuracy, Precision, Recall, F1-score và ROC-AUC. Trong bài toán churn, nhóm ưu tiên "
            "F1-score vì chỉ số này cân bằng giữa Precision và Recall."
        )
        self.add_figure("Hình 5", "So sánh hiệu suất các mô hình", FIG / "bieudocot_sosanh_model.png", width=6.0)
        model_rows = [
            [
                row["model"],
                f"{row['accuracy']:.4f}",
                f"{row['precision']:.4f}",
                f"{row['recall']:.4f}",
                f"{row['f1']:.4f}",
                f"{row['roc_auc']:.4f}",
            ]
            for _, row in results_df.iterrows()
        ]
        self.add_table(
            ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
            model_rows,
            widths=[1.7, 0.8, 0.85, 0.75, 0.65, 0.85],
            font_size=8,
        )
        self.add_table(
            ["Thông tin", "Giá trị"],
            [
                ["Best model", best_summary["best_model"]],
                ["Feature strategy", best_summary["feature_strategy"]],
                ["Accuracy", f"{best_summary['accuracy']:.4f}"],
                ["Precision", f"{best_summary['precision']:.4f}"],
                ["Recall", f"{best_summary['recall']:.4f}"],
                ["F1-score", f"{best_summary['f1']:.4f}"],
                ["ROC-AUC", f"{best_summary['roc_auc']:.4f}"],
            ],
            widths=[2.0, 4.0],
            font_size=10,
        )
        for label, caption, path in self.figures[5:]:
            self.add_figure(label, caption, path, width=5.4 if "Feature" not in caption else 6.0)

        self.doc.add_heading("NHẬN XÉT VÀ Ý NGHĨA THỰC TIỄN", level=1)
        self.add_para(
            f"Mô hình tốt nhất trong lần chạy hiện tại là {best_summary['best_model']}, đạt F1-score "
            f"{best_summary['f1']:.4f} và ROC-AUC {best_summary['roc_auc']:.4f}. ROC-AUC khoảng 0.83 cho thấy mô hình "
            "có khả năng phân biệt tương đối tốt giữa nhóm khách hàng churn và non-churn."
        )
        self.add_para(
            "Trong bài toán churn, Recall và F1-score quan trọng vì doanh nghiệp cần phát hiện càng nhiều khách hàng có nguy cơ rời bỏ càng tốt. "
            "Nếu bỏ sót khách hàng churn, doanh nghiệp có thể mất cơ hội giữ chân khách hàng."
        )
        self.add_para(
            "Các yếu tố quan trọng gồm tenure, Contract, MonthlyCharges, TotalCharges, InternetService, OnlineSecurity, TechSupport và PaymentMethod. "
            "Những yếu tố này gợi ý rằng doanh nghiệp nên chú ý đến khách hàng mới, khách hàng dùng hợp đồng ngắn hạn, khách hàng có chi phí hàng tháng cao "
            "hoặc thiếu dịch vụ hỗ trợ."
        )

        self.doc.add_heading("HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", level=1)
        self.add_bullets(
            [
                "Chưa tuning sâu toàn bộ mô hình bằng GridSearchCV hoặc RandomizedSearchCV.",
                "Chưa sử dụng SHAP hoặc LIME để giải thích từng dự đoán cụ thể.",
                "Dữ liệu không có yếu tố thời gian chi tiết nên chưa triển khai theo hướng time series.",
                "Có thể thử OneHotEncoder thay cho LabelEncoder ở một số mô hình tuyến tính.",
                "Có thể triển khai dashboard hoặc web app để dự đoán churn cho khách hàng mới.",
            ]
        )

        self.doc.add_heading("KẾT LUẬN", level=1)
        self.add_para(
            "Báo cáo đã trình bày đầy đủ quy trình phân tích và dự đoán khách hàng rời bỏ dịch vụ viễn thông. Nhóm đã thực hiện thu thập dữ liệu, "
            "làm sạch dữ liệu, EDA, huấn luyện mô hình, lựa chọn feature, đánh giá mô hình và truyền đạt kết quả."
        )
        self.add_para(
            f"Kết quả tốt nhất thuộc về mô hình {best_summary['best_model']}, với F1-score {best_summary['f1']:.4f} "
            f"và ROC-AUC {best_summary['roc_auc']:.4f}. Mô hình này có thể hỗ trợ doanh nghiệp nhận diện nhóm khách hàng có nguy cơ rời bỏ "
            "để đưa ra chính sách giữ chân phù hợp."
        )
        self.add_para(
            "Về mặt nghiệp vụ, các yếu tố như thời gian sử dụng dịch vụ, loại hợp đồng, chi phí hàng tháng, dịch vụ hỗ trợ và phương thức thanh toán "
            "là những yếu tố đáng chú ý trong việc phân tích churn."
        )

        self.doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
        self.add_numbers(
            [
                "Telco Customer Churn Dataset - Kaggle.",
                "Scikit-learn Documentation: Machine Learning in Python.",
                "XGBoost Documentation.",
                "LightGBM Documentation.",
                "Imbalanced-learn Documentation: SMOTE and SMOTENC.",
                "Tài liệu và hướng dẫn học phần Phân tích dữ liệu với Python.",
            ]
        )

        for section in self.doc.sections:
            footer = section.footer.paragraphs[0]
            footer.text = "Customer Churn Prediction - Báo cáo bài tập lớn"
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(89, 89, 89)

        props = self.doc.core_properties
        props.title = "Phân tích và dự đoán khách hàng rời bỏ dịch vụ viễn thông"
        props.subject = "Customer Churn Prediction"
        props.author = "Customer Churn Analysis Team"
        props.comments = "Generated report for data analysis coursework."

        OUT.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(OUT)
        return OUT


if __name__ == "__main__":
    path = ReportBuilder().build()
    print(path)
